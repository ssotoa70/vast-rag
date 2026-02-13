"""DOCX document parser."""
from pathlib import Path
from docx import Document
from vast_rag.parsers.base import BaseParser
from vast_rag.types import ParsedDocument


class DOCXParser(BaseParser):
    """Parser for Microsoft Word DOCX documents."""

    def can_parse(self, file_path: Path) -> bool:
        """Check if file is DOCX."""
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX and extract text with heading structure."""
        if not self.can_parse(file_path):
            raise ValueError(f"Cannot parse {file_path}: not a DOCX file")

        doc = Document(str(file_path))

        paragraphs = []
        sections = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph is a heading
            if para.style.name.startswith("Heading"):
                sections.append(text)

            paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        metadata = {
            "sections": sections,
            "pages": 1,  # DOCX doesn't expose page count easily
            "format": "docx",
        }

        return ParsedDocument(
            text=full_text,
            metadata=metadata,
            format="docx",
            source_path=file_path,
        )
