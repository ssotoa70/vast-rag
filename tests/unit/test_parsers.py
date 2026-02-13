from pathlib import Path
import pytest
from vast_rag.parsers.markdown import MarkdownParser
from vast_rag.parsers.pdf import PDFParser
from vast_rag.parsers.html import HTMLParser
from vast_rag.parsers.docx import DOCXParser
from vast_rag.parsers.factory import ParserFactory
from vast_rag.types import ParsedDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document


@pytest.fixture
def sample_markdown(tmp_path):
    """Create a sample markdown file."""
    md_file = tmp_path / "sample.md"
    content = """# Main Title

## Section 1

This is section 1 content.

## Section 2

This is section 2 content.

### Subsection 2.1

Nested content here.
"""
    md_file.write_text(content)
    return md_file


def test_markdown_parser_basic(sample_markdown):
    """Test basic markdown parsing."""
    parser = MarkdownParser()
    doc = parser.parse(sample_markdown)

    assert isinstance(doc, ParsedDocument)
    assert doc.format == "markdown"
    assert doc.source_path == sample_markdown
    assert "Main Title" in doc.text
    assert "Section 1" in doc.text


def test_markdown_parser_extracts_sections(sample_markdown):
    """Test that sections are extracted from headers."""
    parser = MarkdownParser()
    doc = parser.parse(sample_markdown)

    sections = doc.sections
    assert "Main Title" in sections
    assert "Section 1" in sections
    assert "Section 2" in sections
    assert "Subsection 2.1" in sections


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF file."""
    pdf_file = tmp_path / "sample.pdf"

    # Create a simple 2-page PDF
    c = canvas.Canvas(str(pdf_file), pagesize=letter)

    # Page 1
    c.drawString(100, 750, "Page 1 Title")
    c.drawString(100, 700, "This is content on page 1.")
    c.showPage()

    # Page 2
    c.drawString(100, 750, "Page 2 Title")
    c.drawString(100, 700, "This is content on page 2.")
    c.showPage()

    c.save()
    return pdf_file


def test_pdf_parser_extracts_pages(sample_pdf):
    """Test PDF parsing extracts page numbers."""
    parser = PDFParser()
    doc = parser.parse(sample_pdf)

    assert doc.format == "pdf"
    assert doc.pages == 2
    assert "Page 1 Title" in doc.text
    assert "Page 2 Title" in doc.text


def test_pdf_parser_page_mapping(sample_pdf):
    """Test that page numbers are tracked in metadata."""
    parser = PDFParser()
    doc = parser.parse(sample_pdf)

    page_texts = doc.metadata["page_texts"]
    assert len(page_texts) == 2
    assert "Page 1 Title" in page_texts[0]
    assert "Page 2 Title" in page_texts[1]


@pytest.fixture
def sample_html(tmp_path):
    """Create a sample HTML file."""
    html_file = tmp_path / "sample.html"
    content = """
    <!DOCTYPE html>
    <html>
    <head><title>Test Document</title></head>
    <body>
        <h1>Main Heading</h1>
        <p>This is a paragraph.</p>
        <h2>Subheading</h2>
        <p>More content here.</p>
        <script>console.log('ignore this');</script>
    </body>
    </html>
    """
    html_file.write_text(content)
    return html_file


def test_html_parser_strips_tags(sample_html):
    """Test HTML parsing removes tags but keeps text."""
    parser = HTMLParser()
    doc = parser.parse(sample_html)

    assert doc.format == "html"
    assert "Main Heading" in doc.text
    assert "<h1>" not in doc.text
    assert "<script>" not in doc.text
    assert "console.log" not in doc.text  # Script removed


def test_html_parser_extracts_headings(sample_html):
    """Test that HTML headings are extracted."""
    parser = HTMLParser()
    doc = parser.parse(sample_html)

    assert "Main Heading" in doc.sections
    assert "Subheading" in doc.sections


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file."""
    docx_file = tmp_path / "sample.docx"

    doc = Document()
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("First paragraph content.")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("Section 1 content.")

    doc.save(str(docx_file))
    return docx_file


def test_docx_parser_extracts_text(sample_docx):
    """Test DOCX parsing extracts all text."""
    parser = DOCXParser()
    doc = parser.parse(sample_docx)

    assert doc.format == "docx"
    assert "Main Title" in doc.text
    assert "First paragraph content" in doc.text
    assert "Section 1" in doc.text


def test_docx_parser_preserves_headings(sample_docx):
    """Test that DOCX headings are extracted."""
    parser = DOCXParser()
    doc = parser.parse(sample_docx)

    assert "Main Title" in doc.sections
    assert "Section 1" in doc.sections


def test_parser_factory_selects_correct_parser(sample_markdown, sample_pdf, sample_html):
    """Test factory selects appropriate parser for each format."""
    factory = ParserFactory()

    md_parser = factory.get_parser(sample_markdown)
    assert md_parser.__class__.__name__ == "MarkdownParser"

    pdf_parser = factory.get_parser(sample_pdf)
    assert pdf_parser.__class__.__name__ == "PDFParser"

    html_parser = factory.get_parser(sample_html)
    assert html_parser.__class__.__name__ == "HTMLParser"


def test_parser_factory_raises_for_unsupported(tmp_path):
    """Test factory raises error for unsupported formats."""
    factory = ParserFactory()
    unsupported = tmp_path / "file.exe"
    unsupported.write_bytes(b"binary data")

    with pytest.raises(ValueError, match="No parser available"):
        factory.get_parser(unsupported)


def test_parser_factory_parse_document(sample_markdown):
    """Test factory can parse documents directly."""
    factory = ParserFactory()
    doc = factory.parse_document(sample_markdown)

    assert doc.format == "markdown"
    assert "Main Title" in doc.text
